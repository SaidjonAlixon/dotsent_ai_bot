from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import logging

logger = logging.getLogger(__name__)

def create_docx(content: str, filename: str, title: str = "") -> str:
    """DOCX fayl yaratish"""
    try:
        os.makedirs("generated_files", exist_ok=True)
        
        doc = Document()
        
        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                paragraph = doc.add_paragraph(para_text)
                
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = 1.5
                paragraph_format.first_line_indent = Inches(0.5)
        
        filepath = f"generated_files/{filename}.docx"
        doc.save(filepath)
        
        logger.info(f"DOCX fayl yaratildi: {filepath}")
        return filepath
    
    except Exception as e:
        logger.error(f"DOCX yaratishda xatolik: {e}")
        raise
