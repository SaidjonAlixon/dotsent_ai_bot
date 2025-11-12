from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def clean_article_content(content):
    """Maqola matnidagi ortiqcha belgilarni olib tashlash"""
    if not content:
        return content
    
    lines = content.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Bo'sh qatorlarni o'tkazib yuborish
        if not line.strip():
            continue
        
        # "###" belgisini olib tashlash
        if line.strip().startswith('###'):
            continue
        
        # Sarlavhalarni olib tashlash
        if line.strip() in ['KIRISH', 'TADQIQOT USLUBLARI', 'NATIJALAR VA MUHOKAMA', 'XULOSA']:
            continue
        
        # Qalin sarlavha belgisini olib tashlash (**Sarlavha**)
        if line.strip().startswith('**') and line.strip().endswith('**'):
            continue
        
        cleaned_lines.append(line)
    
    # Matnni qayta birlashtirish va ortiqcha bo'shliqlarni olib tashlash
    result = ' '.join(' '.join(cleaned_lines).split())
    
    # Agar matn bo'sh bo'lsa, asl matnni qaytarish
    if not result:
        return ' '.join(content.split())
    
    return result


def create_article_document(sections, user_data, file_path):
    """Ilmiy maqola DOCX yaratish"""
    doc = Document()
    
    # Sahifa sozlamalari
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.left_margin = Cm(2.0)   # 20mm
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # Sahifa raqamlari
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # Standart shrift
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.first_line_indent = Cm(1.25)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for section_data in sections:
        section_type = section_data['type']
        
        if section_type == 'title':
            # Sarlavha va mualliflar (1-sahifa)
            title_data = section_data['content']
            
            # Maqola sarlavhasi
            p = doc.add_paragraph(title_data['title'].upper())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(12)
            
            # Mualliflar
            if title_data.get('authors'):
                for author in title_data['authors']:
                    p = doc.add_paragraph(author.get('name', ''))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.size = Pt(14)
                    p.runs[0].font.name = 'Times New Roman'
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.space_after = Pt(6)
                    
                    if author.get('affiliation'):
                        p = doc.add_paragraph(author['affiliation'])
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.runs[0].font.size = Pt(14)
                        p.runs[0].font.name = 'Times New Roman'
                        p.runs[0].font.italic = True
                        p.paragraph_format.first_line_indent = Cm(0)
                        p.paragraph_format.space_after = Pt(12)
            
            doc.add_page_break()
        
        elif section_type == 'annotations':
            # Annotatsiyalar (3 tilda)
            annotations = section_data['content']
            
            # O'zbek tilida
            if 'uz' in annotations:
                p = doc.add_paragraph('ANNOTATSIYA')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(6)
                
                content = annotations['uz']
                # Kalit so'zlarni ajratish
                if 'Kalit so\'zlar:' in content:
                    parts = content.split('Kalit so\'zlar:')
                    # Annotatsiya matni
                    p = doc.add_paragraph(parts[0].strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    # Kalit so'zlar
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(1.25)
                    run = p.add_run('Kalit so\'zlar: ')
                    run.font.bold = True
                    run.font.size = Pt(14)
                    p.add_run(parts[1].strip())
                else:
                    p = doc.add_paragraph(content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                doc.add_paragraph()
            
            # Ingliz tilida
            if 'en' in annotations:
                p = doc.add_paragraph('ABSTRACT')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(6)
                
                content = annotations['en']
                # Keywords ajratish
                if 'Keywords:' in content:
                    parts = content.split('Keywords:')
                    # Abstract matni
                    p = doc.add_paragraph(parts[0].strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    # Keywords
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(1.25)
                    run = p.add_run('Keywords: ')
                    run.font.bold = True
                    run.font.size = Pt(14)
                    p.add_run(parts[1].strip())
                else:
                    p = doc.add_paragraph(content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                doc.add_paragraph()
            
            # Rus tilida
            if 'ru' in annotations:
                p = doc.add_paragraph('АННОТАЦИЯ')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(6)
                
                content = annotations['ru']
                # Ключевые слова ajratish
                if 'Ключевые слова:' in content:
                    parts = content.split('Ключевые слова:')
                    # Аннотация matni
                    p = doc.add_paragraph(parts[0].strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    # Ключевые слова
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(1.25)
                    run = p.add_run('Ключевые слова: ')
                    run.font.bold = True
                    run.font.size = Pt(14)
                    p.add_run(parts[1].strip())
                else:
                    p = doc.add_paragraph(content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                doc.add_paragraph()
        
        elif section_type == 'kirish':
            p = doc.add_paragraph('KIRISH')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(6)
            
            content = section_data['content']
            # Matnni tozalash va uzluksiz yozish
            clean_text = clean_article_content(content)
            p = doc.add_paragraph(clean_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_page_break()
        
        elif section_type == 'methods':
            p = doc.add_paragraph('TADQIQOT USLUBLARI')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(6)
            
            content = section_data['content']
            # Matnni tozalash va uzluksiz yozish
            clean_text = clean_article_content(content)
            p = doc.add_paragraph(clean_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_page_break()
        
        elif section_type == 'results':
            p = doc.add_paragraph('NATIJALAR VA MUHOKAMA')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(6)
            
            content = section_data['content']
            # Matnni tozalash va uzluksiz yozish
            clean_text = clean_article_content(content)
            p = doc.add_paragraph(clean_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_page_break()
        
        elif section_type == 'conclusion':
            p = doc.add_paragraph('XULOSA')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(6)
            
            content = section_data['content']
            # Matnni tozalash va uzluksiz yozish
            clean_text = clean_article_content(content)
            p = doc.add_paragraph(clean_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_page_break()
        
        elif section_type == 'references':
            p = doc.add_paragraph('FOYDALANILGAN ADABIYOTLAR')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(12)
            
            content = section_data['content']
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.space_after = Pt(6)
    
    doc.save(file_path)
    return file_path
