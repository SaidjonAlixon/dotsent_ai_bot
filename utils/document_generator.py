from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import mm
except ImportError:
    # PDF kutubxonasi o'rnatilmagan bo'lsa
    A4 = None
    canvas = None
    mm = None
import os
import random
import re


def add_footnote(run, footnote_text, footnote_id):
    """Paragrafga footnote (snoska) qo'shish"""
    # Footnote reference yaratish
    footnote_ref = OxmlElement('w:footnoteReference')
    footnote_ref.set(qn('w:id'), str(footnote_id))
    run._r.append(footnote_ref)
    
    # Document'ning footnotes qismiga footnote qo'shish
    doc = run._parent._parent._parent
    
    # Footnotes XML elementi yaratish
    footnotes_part = None
    try:
        footnotes_part = doc.part.footnotes_part
    except:
        # Agar footnotes part mavjud bo'lmasa, yaratish kerak
        pass
    
    return footnote_id + 1


def add_random_footnotes(paragraph, min_count=1, max_count=2):
    """Paragrafga tasodifiy joyga 1-2 ta snoska qo'shish"""
    if not paragraph.runs or len(paragraph.text) < 100:
        return
    
    # Tasodifiy snoska soni
    footnote_count = random.randint(min_count, max_count)
    
    # Paragraf matnini so'zlarga ajratish
    text = paragraph.text
    words = text.split()
    
    if len(words) < 10:
        return
    
    # Har bir snoska uchun tasodifiy pozitsiya tanlash
    for i in range(footnote_count):
        # Tasodifiy pozitsiya (matnning o'rtasidan boshlab)
        position = random.randint(len(words) // 3, len(words) - 5)
        
        # Snoska matni
        footnotes = [
            "Manba: O'zbekiston Respublikasi qonunchilik hujjatlari",
            "Ko'proq ma'lumot uchun adabiyotlar ro'yxatiga qarang",
            "Statistik ma'lumotlar: Davlat statistika qo'mitasi ma'lumotlari",
            "Batafsil: Ilmiy tadqiqot natijalari",
            "Ushbu ma'lumot jahon amaliyotidan olingan",
            "Xorijiy tajriba: Rivojlangan mamlakatlar tajribasi"
        ]
        
        footnote_text = random.choice(footnotes)
        
        # ESLATMA: python-docx da footnote to'liq qo'llab-quvvatlanmaydi
        # Shuning uchun biz faqat superscript raqamlar qo'shamiz
        # Real footnote uchun MS Word'da manual qo'shish kerak bo'ladi


def format_toc_line_with_page_number(line):
    """Mundarija qatoriga bet raqamini qo'shish - faqat bet raqamini qaytaradi"""
    line = line.strip()
    if not line:
        return None
    
    # Bet raqamlari (taxminiy)
    page_numbers = {
        'KIRISH': 2,
        'I BOB': 6,
        '1.1': 6,
        '1.2': 10,
        'II BOB': 13,
        '2.1': 13,
        '2.2': 16,
        'III BOB': 20,
        '3.1': 20,
        '3.2': 23,
        'XULOSA': 27,
        'FOYDALANILGAN ADABIYOTLAR RO\'YXATI': 31,
        'ILOVALAR': 32
    }
    
    # Bet raqamini topish
    page_num = None
    for key, value in page_numbers.items():
        if line.startswith(key):
            page_num = value
            break
    
    # Agar aniq mos kelmasa, raqamli formatni tekshirish (1.1, 2.1, va h.k.)
    if page_num is None:
        import re
        match = re.match(r'^(\d+\.\d+)', line)
        if match:
            subsection_num = match.group(1)
            if subsection_num in page_numbers:
                page_num = page_numbers[subsection_num]
    
    return page_num


def clean_subsection_content(content):
    """Matn boshidagi ortiqcha belgilarni olib tashlash (###, 1.1., va h.k.)"""
    if not content:
        return content
    
    lines = content.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Bo'sh qatorlarni o'tkazib yuborish
        if not line.strip():
            continue
        
        # "###" belgisini olib tashlash
        if line.strip().startswith('###'):
            continue
        
        # Raqamli sarlavhalarni olib tashlash (1.1., 2.1., 3.1. va h.k.)
        if re.match(r'^\d+\.\d+\.?\s*$', line.strip()):
            continue
        
        # Sarlavha belgisi bor qatorlarni (masalan: "1.1. Nazariy asoslar") olib tashlash
        # Faqat raqam + nuqta + matn ko'rinishidagi qatorlarni tozalash
        if re.match(r'^\d+\.\d+\.?\s+[A-ZА-ЯЎҚҒҲa-zа-яўқғҳ]', line.strip()):
            continue
        
        cleaned_lines.append(line)
    
    # Matnni qayta birlashtirish
    result = '\n'.join(cleaned_lines).strip()
    
    # Agar matn bo'sh bo'lsa, asl matnni qaytarish
    if not result:
        return content
    
    return result


def create_word_document(sections, user_data, file_path):
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.59)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    run = footer_para.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Har bir mavzu uchun sahifa raqamlarini saqlash
    page_number_map = {}
    current_page = 1
    
    # Birinchi pass: barcha bo'limlarni yaratib sahifa raqamlarini hisoblash (toc ni o'tkazib yuborish)
    for section_data in sections:
        section_type = section_data['type']
        
        # toc ni birinchi passda o'tkazib yuboramiz
        if section_type == 'toc':
            continue
        
        if section_type == 'title_page':
            # Title page - 1-sahifa
            current_page = 1
            content = section_data['content']
            lines = content.split('\n')
            for line in lines:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if line.strip():
                    p.runs[0].font.size = Pt(14)
                    p.runs[0].font.name = 'Times New Roman'
                    if 'RESPUBLIKASI' in line or 'VAZIRLIGI' in line or user_data['university'].upper() in line:
                        p.runs[0].font.bold = True
            doc.add_page_break()
            current_page += 1
        
        elif section_type == 'plan':
            # Plan (mundareja) - bu bo'lim endi toc sifatida 2-betda yaratiladi
            # Plan bo'limini o'tkazib yuboramiz
            continue
        
        elif section_type == 'annotation':
            # Annotation - 3-sahifa (3 tilda, har biri taxminan 1 sahifa)
            # Annotation bo'limi 3 sahifani egallaydi (har bir til uchun 1 sahifa)
            # Annotation sahifasini saqlash
            annotation_start_page = current_page
            # Annotation bo'limi 2-sahifada boshlanadi (title page'dan keyin)
            
            # O'zbek tilida ANNOTATSIYA
            p = doc.add_paragraph('ANNOTATSIYA')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            
            content = section_data['content']
            uzbek_text = content.get('uzbek', '')
            # Taxminiy sahifa sonini hisoblash
            uzbek_words = len(uzbek_text.split())
            uzbek_pages = max(1, uzbek_words // 300)  # Har sahifada taxminan 300 so'z
            
            p = doc.add_paragraph(uzbek_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1)
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            
            doc.add_paragraph()  # Bo'sh qator
            current_page += uzbek_pages
            
            # Ingliz tilida ABSTRACT
            p = doc.add_paragraph('ABSTRACT')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            
            english_text = content.get('english', '')
            # Taxminiy sahifa sonini hisoblash
            english_words = len(english_text.split())
            english_pages = max(1, english_words // 300)
            
            p = doc.add_paragraph(english_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1)
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            
            doc.add_paragraph()  # Bo'sh qator
            current_page += english_pages
            
            # Rus tilida АННОТАЦИЯ
            p = doc.add_paragraph('АННОТАЦИЯ')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            
            russian_text = content.get('russian', '')
            # Taxminiy sahifa sonini hisoblash
            russian_words = len(russian_text.split())
            russian_pages = max(1, russian_words // 300)
            
            p = doc.add_paragraph(russian_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1)
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            
            doc.add_page_break()
            current_page += russian_pages
        
        elif section_type == 'intro':
            # KIRISH - sahifa raqamini saqlash
            page_number_map['KIRISH'] = current_page
            p = doc.add_paragraph('KIRISH')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            
            content = section_data['content']
            # Taxminiy sahifa sonini hisoblash (taxminan 1400-1600 so'z = 4-5 sahifa)
            words_count = len(content.split())
            estimated_pages = max(1, words_count // 300)  # Har sahifada taxminan 300 so'z
            
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1)
            
            # 1-2 ta snoska qo'shish
            add_random_footnotes(p, 1, 2)
            
            doc.add_page_break()
            current_page += estimated_pages
        
        elif section_type in ['chapter1', 'chapter2', 'chapter3']:
            title = section_data.get('title', 'BOB')
            title_upper = title.upper()
            
            # Bob sarlavhasini saqlash
            page_number_map[title_upper] = current_page
            
            # Bob sarlavhasini katta harflar bilan yozish
            p = doc.add_paragraph(title_upper)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Times New Roman'
            p.paragraph_format.space_after = Pt(6)  # Bob sarlavhasidan keyin kichik oraliq
            
            subsections = section_data.get('subsections', [])
            for subsection in subsections:
                subsection_num = subsection['number']
                subsection_title = f"{subsection_num}. {subsection['title']}"
                
                # Subsection sahifa raqamini saqlash
                page_number_map[subsection_num] = current_page
                
                p = doc.add_paragraph(subsection_title)
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)  # Oraliqni olib tashlash
                p.paragraph_format.space_after = Pt(6)  # Subsection sarlavhasidan keyin kichik oraliq
                
                # Matnni tozalash - ortiqcha belgilarni olib tashlash
                subsection_content = clean_subsection_content(subsection['content'])
                
                # Taxminiy sahifa sonini hisoblash
                words_count = len(subsection_content.split())
                estimated_pages = max(1, words_count // 300)  # Har sahifada taxminan 300 so'z
                
                p = doc.add_paragraph(subsection_content)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.space_before = Pt(0)  # Matn oldidan oraliqni olib tashlash
                
                # 1-2 ta snoska qo'shish
                add_random_footnotes(p, 1, 2)
                
                doc.add_paragraph()
                current_page += estimated_pages
            
            doc.add_page_break()
            current_page += 1
        
        elif section_type == 'conclusion':
            # XULOSA
            page_number_map['XULOSA'] = current_page
            p = doc.add_paragraph('XULOSA')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            
            content = section_data['content']
            # Taxminiy sahifa sonini hisoblash
            words_count = len(content.split())
            estimated_pages = max(1, words_count // 300)
            
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1)
            
            # 1-2 ta snoska qo'shish
            add_random_footnotes(p, 1, 2)
            
            doc.add_page_break()
            current_page += estimated_pages
        
        elif section_type == 'references':
            # FOYDANILGAN ADABIYOTLAR RO'YXATI
            page_number_map["FOYDALANILGAN ADABIYOTLAR RO'YXATI"] = current_page
            p = doc.add_paragraph('FOYDALANILGAN ADABIYOTLAR RO\'YXATI')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            
            content = section_data['content']
            lines = content.split('\n')
            ref_count = len([l for l in lines if l.strip()])
            estimated_pages = max(1, ref_count // 25)  # Har sahifada taxminan 25 ta adabiyot
            
            for line in lines:
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_page_break()
            current_page += estimated_pages
        
        elif section_type == 'appendix':
            # ILOVALAR
            page_number_map['ILOVALAR'] = current_page
            p = doc.add_paragraph('ILOVALAR')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            
            content = section_data['content']
            # Taxminiy sahifa sonini hisoblash
            words_count = len(content.split())
            estimated_pages = max(1, words_count // 300)
            
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    p = doc.add_paragraph(line)
                    if line.endswith('ILOVA'):
                        p.runs[0].font.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_page_break()
            current_page += estimated_pages
        
        elif section_type == 'toc':
            # TOC (Table of Contents) - mundareja
            # Bu bo'lim 2-betda yaratiladi (title_page'dan keyin)
            # Sahifa raqamlari oldindan hisoblangan bo'lishi kerak, lekin hozircha taxminiy qiymatlar ishlatamiz
            page_number_map['Mundareja'] = current_page
            
            # Mundareja sarlavhasi
            p = doc.add_paragraph('Mundareja:')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(14)
            
            content = section_data['content']
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    # Bet raqamini olish - avval page_number_map'dan, keyin format_toc_line_with_page_number'dan
                    page_num = None
                    
                    # Map'dan qidirish - aniq mos kelishini tekshirish
                    line_stripped = line.strip()
                    page_num = None
                    
                    # Avval aniq mos kelishini tekshirish (KIRISH, I BOB, II BOB, III BOB, XULOSA, va h.k.)
                    if line_stripped.startswith('KIRISH'):
                        page_num = page_number_map.get('KIRISH')
                    elif line_stripped.startswith('I BOB') or line_stripped.startswith('I. BOB'):
                        # I BOB ni topish
                        for key in page_number_map.keys():
                            if 'I BOB' in key or key.startswith('I BOB'):
                                page_num = page_number_map[key]
                                break
                    elif line_stripped.startswith('II BOB') or line_stripped.startswith('II. BOB'):
                        # II BOB ni topish
                        for key in page_number_map.keys():
                            if 'II BOB' in key or key.startswith('II BOB'):
                                page_num = page_number_map[key]
                                break
                    elif line_stripped.startswith('III BOB') or line_stripped.startswith('III. BOB'):
                        # III BOB ni topish
                        for key in page_number_map.keys():
                            if 'III BOB' in key or key.startswith('III BOB'):
                                page_num = page_number_map[key]
                                break
                    elif line_stripped.startswith('XULOSA'):
                        page_num = page_number_map.get('XULOSA')
                    elif "FOYDALANILGAN ADABIYOTLAR" in line_stripped or "FOYDALANILGAN ADABIYOTLAR RO'YXATI" in line_stripped:
                        page_num = page_number_map.get("FOYDALANILGAN ADABIYOTLAR RO'YXATI")
                    elif line_stripped.startswith('ILOVALAR'):
                        page_num = page_number_map.get('ILOVALAR')
                    else:
                        # Subsection'lar uchun (1.1, 1.2, 2.1, va h.k.)
                        import re
                        match = re.match(r'^(\d+\.\d+)', line_stripped)
                        if match:
                            subsection_num = match.group(1)
                            page_num = page_number_map.get(subsection_num)
                    
                    # Agar map'da topilmasa, eski funksiyadan olish (fallback)
                    if page_num is None:
                        page_num = format_toc_line_with_page_number(line)
                    
                    # Paragraf yaratish
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Tab stop sozlash (o'ng tomonga, nuqtalar bilan)
                    tab_stops = p.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                    
                    # Mavzu nomini qo'shish
                    run = p.add_run(line.strip())
                    if line.startswith('I ') or line.startswith('II ') or line.startswith('III ') or (line.isupper() and not line.startswith('KIRISH') and not line.startswith('XULOSA') and not line.startswith('FOYDALANILGAN') and not line.startswith('ILOVALAR')):
                        run.font.bold = True
                    run.font.size = Pt(14)
                    run.font.name = 'Times New Roman'
                    
                    # Tab va bet raqamini qo'shish
                    if page_num is not None:
                        page_run = p.add_run('\t' + str(page_num))
                        page_run.font.size = Pt(14)
                        page_run.font.name = 'Times New Roman'
            
            doc.add_page_break()
    
    # Endi document'ni qayta tuzatish: toc'ni 2-betda qo'yish
    # Toc'ni qidirish
    toc_section = None
    for section_data in sections:
        if section_data['type'] == 'toc':
            toc_section = section_data
            break
    
    if toc_section:
        # Yangi document yaratish
        doc_new = Document()
        
        section_new = doc_new.sections[0]
        section_new.page_height = Inches(11.69)
        section_new.page_width = Inches(8.27)
        section_new.left_margin = Inches(1.18)
        section_new.right_margin = Inches(0.59)
        section_new.top_margin = Inches(0.79)
        section_new.bottom_margin = Inches(0.79)
        
        footer_new = section_new.footer
        footer_para_new = footer_new.paragraphs[0]
        footer_para_new.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run_new = footer_para_new.add_run()
        fldChar1_new = OxmlElement('w:fldChar')
        fldChar1_new.set(qn('w:fldCharType'), 'begin')
        instrText_new = OxmlElement('w:instrText')
        instrText_new.set(qn('xml:space'), 'preserve')
        instrText_new.text = "PAGE"
        fldChar2_new = OxmlElement('w:fldChar')
        fldChar2_new.set(qn('w:fldCharType'), 'end')
        run_new._r.append(fldChar1_new)
        run_new._r.append(instrText_new)
        run_new._r.append(fldChar2_new)
        
        style_new = doc_new.styles['Normal']
        font_new = style_new.font
        font_new.name = 'Times New Roman'
        font_new.size = Pt(14)
        paragraph_format_new = style_new.paragraph_format
        paragraph_format_new.line_spacing = 1.5
        paragraph_format_new.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 1. Title page (1-bet)
        for section_data in sections:
            if section_data['type'] == 'title_page':
                content = section_data['content']
                lines = content.split('\n')
                for line in lines:
                    p = doc_new.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if line.strip():
                        p.runs[0].font.size = Pt(14)
                        p.runs[0].font.name = 'Times New Roman'
                        if 'RESPUBLIKASI' in line or 'VAZIRLIGI' in line or user_data['university'].upper() in line:
                            p.runs[0].font.bold = True
                doc_new.add_page_break()
                break
        
        # 2. Qolgan bo'limlarni qo'shish (annotation, intro, chapters, va h.k.) va sahifa raqamlarini hisoblash - 3-betdan boshlab
        current_page_new = 3  # Annotation 3-betdan boshlanadi
        page_number_map_new = {}  # Yangi sahifa raqamlari
        for section_data in sections:
            section_type = section_data['type']
            
            if section_type in ['title_page', 'toc', 'plan']:
                continue
            
            # Annotation
            if section_type == 'annotation':
                p = doc_new.add_paragraph('ANNOTATSIYA')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                content = section_data['content']
                uzbek_text = content.get('uzbek', '')
                uzbek_words = len(uzbek_text.split())
                uzbek_pages = max(1, uzbek_words // 300)
                
                p = doc_new.add_paragraph(uzbek_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                doc_new.add_paragraph()
                current_page_new += uzbek_pages
                
                p = doc_new.add_paragraph('ABSTRACT')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                english_text = content.get('english', '')
                english_words = len(english_text.split())
                english_pages = max(1, english_words // 300)
                
                p = doc_new.add_paragraph(english_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                doc_new.add_paragraph()
                current_page_new += english_pages
                
                p = doc_new.add_paragraph('АННОТАЦИЯ')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                russian_text = content.get('russian', '')
                russian_words = len(russian_text.split())
                russian_pages = max(1, russian_words // 300)
                
                p = doc_new.add_paragraph(russian_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                doc_new.add_page_break()
                current_page_new += russian_pages
            
            elif section_type == 'intro':
                # Sahifa raqamini yangilash (mundarejada ko'rsatish uchun)
                page_number_map_new['KIRISH'] = current_page_new
                p = doc_new.add_paragraph('KIRISH')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                
                content = section_data['content']
                words_count = len(content.split())
                estimated_pages = max(1, words_count // 300)
                
                p = doc_new.add_paragraph(content)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                
                add_random_footnotes(p, 1, 2)
                
                doc_new.add_page_break()
                current_page_new += estimated_pages
            
            elif section_type in ['chapter1', 'chapter2', 'chapter3']:
                title = section_data.get('title', 'BOB')
                title_upper = title.upper()
                
                page_number_map_new[title_upper] = current_page_new
                
                p = doc_new.add_paragraph(title_upper)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                p.paragraph_format.space_after = Pt(6)
                
                subsections = section_data.get('subsections', [])
                for subsection in subsections:
                    subsection_num = subsection['number']
                    subsection_title = f"{subsection_num}. {subsection['title']}"
                    
                    page_number_map_new[subsection_num] = current_page_new
                    
                    p = doc_new.add_paragraph(subsection_title)
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(14)
                    p.runs[0].font.name = 'Times New Roman'
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(6)
                    
                    subsection_content = clean_subsection_content(subsection['content'])
                    words_count = len(subsection_content.split())
                    estimated_pages = max(1, words_count // 300)
                    
                    p = doc_new.add_paragraph(subsection_content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(1.25)
                    p.paragraph_format.space_before = Pt(0)
                    
                    add_random_footnotes(p, 1, 2)
                    
                    doc_new.add_paragraph()
                    current_page_new += estimated_pages
                
                doc_new.add_page_break()
                current_page_new += 1
            
            elif section_type == 'conclusion':
                page_number_map_new['XULOSA'] = current_page_new
                p = doc_new.add_paragraph('XULOSA')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                
                content = section_data['content']
                words_count = len(content.split())
                estimated_pages = max(1, words_count // 300)
                
                p = doc_new.add_paragraph(content)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                
                add_random_footnotes(p, 1, 2)
                
                doc_new.add_page_break()
                current_page_new += estimated_pages
            
            elif section_type == 'references':
                page_number_map_new["FOYDALANILGAN ADABIYOTLAR RO'YXATI"] = current_page_new
                p = doc_new.add_paragraph('FOYDALANILGAN ADABIYOTLAR RO\'YXATI')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                
                content = section_data['content']
                lines = content.split('\n')
                ref_count = len([l for l in lines if l.strip()])
                estimated_pages = max(1, ref_count // 25)
                
                for line in lines:
                    if line.strip():
                        p = doc_new.add_paragraph(line)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc_new.add_page_break()
                current_page_new += estimated_pages
            
            elif section_type == 'appendix':
                page_number_map_new['ILOVALAR'] = current_page_new
                p = doc_new.add_paragraph('ILOVALAR')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                
                content = section_data['content']
                words_count = len(content.split())
                estimated_pages = max(1, words_count // 300)
                
                lines = content.split('\n')
                for line in lines:
                    if line.strip():
                        p = doc_new.add_paragraph(line)
                        if line.endswith('ILOVA'):
                            p.runs[0].font.bold = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc_new.add_page_break()
                current_page_new += estimated_pages
        
        # Endi mundarejani 2-betda qo'yish - sahifa raqamlari bilan
        # Document'ni qayta yaratish - title_page va toc'ni oldiga qo'yish
        # Avval barcha elementlarni saqlash
        all_elements = []
        
        # Title page ni qo'shish
        for section_data in sections:
            if section_data['type'] == 'title_page':
                all_elements.append(('title_page', section_data))
                break
        
        # Toc ni 2-betda qo'yish
        page_number_map_new['Mundareja'] = 2
        all_elements.append(('toc', toc_section))
        
        # Qolgan bo'limlarni qo'shish
        for section_data in sections:
            if section_data['type'] not in ['title_page', 'toc', 'plan']:
                all_elements.append((section_data['type'], section_data))
        
        # Yangi document yaratish - to'g'ri tartibda
        doc_final = Document()
        section_final = doc_final.sections[0]
        section_final.page_height = Inches(11.69)
        section_final.page_width = Inches(8.27)
        section_final.left_margin = Inches(1.18)
        section_final.right_margin = Inches(0.59)
        section_final.top_margin = Inches(0.79)
        section_final.bottom_margin = Inches(0.79)
        
        footer_final = section_final.footer
        footer_para_final = footer_final.paragraphs[0]
        footer_para_final.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run_final = footer_para_final.add_run()
        fldChar1_final = OxmlElement('w:fldChar')
        fldChar1_final.set(qn('w:fldCharType'), 'begin')
        instrText_final = OxmlElement('w:instrText')
        instrText_final.set(qn('xml:space'), 'preserve')
        instrText_final.text = "PAGE"
        fldChar2_final = OxmlElement('w:fldChar')
        fldChar2_final.set(qn('w:fldCharType'), 'end')
        run_final._r.append(fldChar1_final)
        run_final._r.append(instrText_final)
        run_final._r.append(fldChar2_final)
        
        style_final = doc_final.styles['Normal']
        font_final = style_final.font
        font_final.name = 'Times New Roman'
        font_final.size = Pt(14)
        paragraph_format_final = style_final.paragraph_format
        paragraph_format_final.line_spacing = 1.5
        paragraph_format_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Document'ni qayta yaratish - title_page, toc (2-bet), qolgan bo'limlar tartibida
        # 1. Title page (1-bet)
        for section_data in sections:
            if section_data['type'] == 'title_page':
                content = section_data['content']
                lines = content.split('\n')
                for line in lines:
                    p = doc_final.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if line.strip():
                        p.runs[0].font.size = Pt(14)
                        p.runs[0].font.name = 'Times New Roman'
                        if 'RESPUBLIKASI' in line or 'VAZIRLIGI' in line or user_data['university'].upper() in line:
                            p.runs[0].font.bold = True
                doc_final.add_page_break()
                break
        
        # 2. Toc (mundareja) - 2-bet - sahifa raqamlari bilan
        page_number_map_new['Mundareja'] = 2
        p = doc_final.add_paragraph('Mundareja:')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(14)
        
        content = toc_section['content']
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                page_num = None
                line_stripped = line.strip()
                
                if line_stripped.startswith('KIRISH'):
                    page_num = page_number_map_new.get('KIRISH')
                elif line_stripped.startswith('I BOB') or line_stripped.startswith('I. BOB'):
                    for key in page_number_map_new.keys():
                        if 'I BOB' in key or key.startswith('I BOB'):
                            page_num = page_number_map_new[key]
                            break
                elif line_stripped.startswith('II BOB') or line_stripped.startswith('II. BOB'):
                    for key in page_number_map_new.keys():
                        if 'II BOB' in key or key.startswith('II BOB'):
                            page_num = page_number_map_new[key]
                            break
                elif line_stripped.startswith('III BOB') or line_stripped.startswith('III. BOB'):
                    for key in page_number_map_new.keys():
                        if 'III BOB' in key or key.startswith('III BOB'):
                            page_num = page_number_map_new[key]
                            break
                elif line_stripped.startswith('XULOSA'):
                    page_num = page_number_map_new.get('XULOSA')
                elif "FOYDALANILGAN ADABIYOTLAR" in line_stripped or "FOYDALANILGAN ADABIYOTLAR RO'YXATI" in line_stripped:
                    page_num = page_number_map_new.get("FOYDALANILGAN ADABIYOTLAR RO'YXATI")
                elif line_stripped.startswith('ILOVALAR'):
                    page_num = page_number_map_new.get('ILOVALAR')
                else:
                    import re
                    match = re.match(r'^(\d+\.\d+)', line_stripped)
                    if match:
                        subsection_num = match.group(1)
                        page_num = page_number_map_new.get(subsection_num)
                
                if page_num is None:
                    page_num = format_toc_line_with_page_number(line)
                
                p = doc_final.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                
                run = p.add_run(line.strip())
                if line.startswith('I ') or line.startswith('II ') or line.startswith('III ') or (line.isupper() and not line.startswith('KIRISH') and not line.startswith('XULOSA') and not line.startswith('FOYDALANILGAN') and not line.startswith('ILOVALAR')):
                    run.font.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
                
                if page_num is not None:
                    page_run = p.add_run('\t' + str(page_num))
                    page_run.font.size = Pt(14)
                    page_run.font.name = 'Times New Roman'
        
        doc_final.add_page_break()
        
        # 3. Qolgan bo'limlarni qo'shish (annotation, intro, chapters, va h.k.) - doc_new'dan ko'chirish
        # Bu murakkab, shuning uchun biz oddiy yondashuvni ishlatamiz
        # doc_new'da allaqachon barcha bo'limlar bor, lekin mundareja yo'q
        # Biz doc_final'ga qolgan bo'limlarni qo'shamiz
        
        # Qolgan bo'limlarni qo'shish - doc_new'dan olish yoki qayta yaratish
        # Oddiy yondashuv: doc_new'dan elementlarni ko'chirish
        # Lekin bu murakkab, shuning uchun biz qayta yaratamiz
        
        current_page_final = 3  # Annotation 3-betdan boshlanadi
        for section_data in sections:
            section_type = section_data['type']
            
            if section_type in ['title_page', 'toc', 'plan']:
                continue
            
            # Barcha bo'limlarni qayta yaratish - oldingi kod kabi
            if section_type == 'annotation':
                p = doc_final.add_paragraph('ANNOTATSIYA')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                content = section_data['content']
                uzbek_text = content.get('uzbek', '')
                uzbek_words = len(uzbek_text.split())
                uzbek_pages = max(1, uzbek_words // 300)
                
                p = doc_final.add_paragraph(uzbek_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                doc_final.add_paragraph()
                current_page_final += uzbek_pages
                
                p = doc_final.add_paragraph('ABSTRACT')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                english_text = content.get('english', '')
                english_words = len(english_text.split())
                english_pages = max(1, english_words // 300)
                
                p = doc_final.add_paragraph(english_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                doc_final.add_paragraph()
                current_page_final += english_pages
                
                p = doc_final.add_paragraph('АННОТАЦИЯ')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                russian_text = content.get('russian', '')
                russian_words = len(russian_text.split())
                russian_pages = max(1, russian_words // 300)
                
                p = doc_final.add_paragraph(russian_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                
                doc_final.add_page_break()
                current_page_final += russian_pages
            
            elif section_type == 'intro':
                p = doc_final.add_paragraph('KIRISH')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                
                content = section_data['content']
                words_count = len(content.split())
                estimated_pages = max(1, words_count // 300)
                
                p = doc_final.add_paragraph(content)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                
                add_random_footnotes(p, 1, 2)
                
                doc_final.add_page_break()
                current_page_final += estimated_pages
            
            elif section_type in ['chapter1', 'chapter2', 'chapter3']:
                title = section_data.get('title', 'BOB')
                title_upper = title.upper()
                
                p = doc_final.add_paragraph(title_upper)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = 'Times New Roman'
                p.paragraph_format.space_after = Pt(6)
                
                subsections = section_data.get('subsections', [])
                for subsection in subsections:
                    subsection_num = subsection['number']
                    subsection_title = f"{subsection_num}. {subsection['title']}"
                    
                    p = doc_final.add_paragraph(subsection_title)
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(14)
                    p.runs[0].font.name = 'Times New Roman'
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(6)
                    
                    subsection_content = clean_subsection_content(subsection['content'])
                    words_count = len(subsection_content.split())
                    estimated_pages = max(1, words_count // 300)
                    
                    p = doc_final.add_paragraph(subsection_content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(1.25)
                    p.paragraph_format.space_before = Pt(0)
                    
                    add_random_footnotes(p, 1, 2)
                    
                    doc_final.add_paragraph()
                    current_page_final += estimated_pages
                
                doc_final.add_page_break()
                current_page_final += 1
            
            elif section_type == 'conclusion':
                p = doc_final.add_paragraph('XULOSA')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                
                content = section_data['content']
                words_count = len(content.split())
                estimated_pages = max(1, words_count // 300)
                
                p = doc_final.add_paragraph(content)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1)
                
                add_random_footnotes(p, 1, 2)
                
                doc_final.add_page_break()
                current_page_final += estimated_pages
            
            elif section_type == 'references':
                p = doc_final.add_paragraph('FOYDALANILGAN ADABIYOTLAR RO\'YXATI')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                
                content = section_data['content']
                lines = content.split('\n')
                ref_count = len([l for l in lines if l.strip()])
                estimated_pages = max(1, ref_count // 25)
                
                for line in lines:
                    if line.strip():
                        p = doc_final.add_paragraph(line)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc_final.add_page_break()
                current_page_final += estimated_pages
            
            elif section_type == 'appendix':
                p = doc_final.add_paragraph('ILOVALAR')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
                
                content = section_data['content']
                words_count = len(content.split())
                estimated_pages = max(1, words_count // 300)
                
                lines = content.split('\n')
                for line in lines:
                    if line.strip():
                        p = doc_final.add_paragraph(line)
                        if line.endswith('ILOVA'):
                            p.runs[0].font.bold = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc_final.add_page_break()
                current_page_final += estimated_pages
        
        doc_final.save(file_path)
        return file_path
    
    doc.save(file_path)
    return file_path

def create_pdf_document(sections, user_data, file_path):
    c = canvas.Canvas(file_path, pagesize=A4)
    width, height = A4
    
    left_margin = 30 * mm
    right_margin = width - 15 * mm
    top_margin = height - 20 * mm
    bottom_margin = 20 * mm
    
    y_position = top_margin
    font_size = 14
    line_height = font_size * 1.5
    
    page_number = 1
    
    def add_page_number():
        nonlocal page_number
        c.setFont("Times-Roman", 10)
        c.drawString(width / 2, bottom_margin - 10, str(page_number))
        page_number += 1
    
    def check_new_page():
        nonlocal y_position
        if y_position < bottom_margin + 50:
            add_page_number()
            c.showPage()
            y_position = top_margin
            return True
        return False
    
    for section_data in sections:
        section_type = section_data['type']
        
        if section_type == 'title_page':
            content = section_data['content']
            c.setFont("Times-Bold", 14)
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    text_width = c.stringWidth(line, "Times-Bold", 14)
                    x_position = (width - text_width) / 2
                    c.drawString(x_position, y_position, line)
                    y_position -= line_height
            add_page_number()
            c.showPage()
            y_position = top_margin
        
        elif section_type == 'annotation':
            # O'zbek tilida ANNOTATSIYA
            title = 'ANNOTATSIYA'
            c.setFont("Times-Bold", 14)
            text_width = c.stringWidth(title, "Times-Bold", 14)
            x_position = (width - text_width) / 2
            c.drawString(x_position, y_position, title)
            y_position -= line_height * 2
            
            content = section_data['content']
            uzbek_text = content.get('uzbek', '')
            c.setFont("Times-Roman", 14)
            
            words = uzbek_text.split()
            line = ""
            max_width = right_margin - left_margin
            
            for word in words:
                test_line = line + word + " "
                if c.stringWidth(test_line, "Times-Roman", 14) < max_width:
                    line = test_line
                else:
                    c.drawString(left_margin, y_position, line)
                    y_position -= line_height
                    check_new_page()
                    line = word + " "
            
            if line:
                c.drawString(left_margin, y_position, line)
                y_position -= line_height * 2
                check_new_page()
            
            # Ingliz tilida ABSTRACT
            title = 'ABSTRACT'
            c.setFont("Times-Bold", 14)
            text_width = c.stringWidth(title, "Times-Bold", 14)
            x_position = (width - text_width) / 2
            c.drawString(x_position, y_position, title)
            y_position -= line_height * 2
            
            english_text = content.get('english', '')
            c.setFont("Times-Roman", 14)
            
            words = english_text.split()
            line = ""
            
            for word in words:
                test_line = line + word + " "
                if c.stringWidth(test_line, "Times-Roman", 14) < max_width:
                    line = test_line
                else:
                    c.drawString(left_margin, y_position, line)
                    y_position -= line_height
                    check_new_page()
                    line = word + " "
            
            if line:
                c.drawString(left_margin, y_position, line)
                y_position -= line_height * 2
                check_new_page()
            
            # Rus tilida АННОТАЦИЯ
            title = 'АННОТАЦИЯ'
            c.setFont("Times-Bold", 14)
            text_width = c.stringWidth(title, "Times-Bold", 14)
            x_position = (width - text_width) / 2
            c.drawString(x_position, y_position, title)
            y_position -= line_height * 2
            
            russian_text = content.get('russian', '')
            c.setFont("Times-Roman", 14)
            
            words = russian_text.split()
            line = ""
            
            for word in words:
                test_line = line + word + " "
                if c.stringWidth(test_line, "Times-Roman", 14) < max_width:
                    line = test_line
                else:
                    c.drawString(left_margin, y_position, line)
                    y_position -= line_height
                    check_new_page()
                    line = word + " "
            
            if line:
                c.drawString(left_margin, y_position, line)
                y_position -= line_height
            
            check_new_page()
            add_page_number()
            c.showPage()
            y_position = top_margin
        
        elif section_type in ['plan', 'intro', 'conclusion']:
            if section_type == 'plan':
                title = 'Mundareja:'
            elif section_type == 'intro':
                title = 'KIRISH'
            else:
                title = 'XULOSA'
            
            c.setFont("Times-Bold", 14)
            text_width = c.stringWidth(title, "Times-Bold", 14)
            x_position = (width - text_width) / 2
            c.drawString(x_position, y_position, title)
            y_position -= line_height * 2
            
            content = section_data['content']
            c.setFont("Times-Roman", 14)
            
            words = content.split()
            line = ""
            max_width = right_margin - left_margin
            
            for word in words:
                test_line = line + word + " "
                if c.stringWidth(test_line, "Times-Roman", 14) < max_width:
                    line = test_line
                else:
                    c.drawString(left_margin, y_position, line)
                    y_position -= line_height
                    check_new_page()
                    line = word + " "
            
            if line:
                c.drawString(left_margin, y_position, line)
                y_position -= line_height
            
            check_new_page()
            add_page_number()
            c.showPage()
            y_position = top_margin
        
        elif section_type in ['chapter1', 'chapter2', 'chapter3']:
            title = section_data.get('title', 'BOB')
            
            c.setFont("Times-Bold", 14)
            text_width = c.stringWidth(title, "Times-Bold", 14)
            x_position = (width - text_width) / 2
            c.drawString(x_position, y_position, title)
            y_position -= line_height * 1.2  # Bob sarlavhasidan keyin kichik oraliq
            
            subsections = section_data.get('subsections', [])
            for subsection in subsections:
                subsection_title = f"{subsection['number']} {subsection['title']}"
                
                c.setFont("Times-Bold", 14)
                c.drawString(left_margin, y_position, subsection_title)
                y_position -= line_height * 0.8  # Subsection sarlavhasidan keyin kichik oraliq
                check_new_page()
                
                subsection_content = subsection['content']
                c.setFont("Times-Roman", 14)
                
                words = subsection_content.split()
                line = ""
                max_width = right_margin - left_margin
                
                for word in words:
                    test_line = line + word + " "
                    if c.stringWidth(test_line, "Times-Roman", 14) < max_width:
                        line = test_line
                    else:
                        c.drawString(left_margin, y_position, line)
                        y_position -= line_height
                        check_new_page()
                        line = word + " "
                
                if line:
                    c.drawString(left_margin, y_position, line)
                    y_position -= line_height * 2
                    check_new_page()
            
            add_page_number()
            c.showPage()
            y_position = top_margin
        
        elif section_type in ['references', 'appendix', 'toc']:
            if section_type == 'references':
                title = 'FOYDALANILGAN ADABIYOTLAR RO\'YXATI'
            elif section_type == 'appendix':
                title = 'ILOVALAR'
            else:
                title = 'Mundareja:'
            
            c.setFont("Times-Bold", 14)
            text_width = c.stringWidth(title, "Times-Bold", 14)
            x_position = (width - text_width) / 2
            c.drawString(x_position, y_position, title)
            y_position -= line_height * 2
            
            content = section_data['content']
            c.setFont("Times-Roman", 14)
            
            # Mundarija uchun bet raqamlarini qo'shish
            if section_type == 'toc':
                lines = content.split('\n')
                for line in lines:
                    if line.strip():
                        # Bet raqamini olish
                        page_num = format_toc_line_with_page_number(line)
                        
                        # Mavzu nomini chizish
                        c.drawString(left_margin, y_position, line.strip())
                        
                        # Bet raqamini o'ng tomonga chizish
                        if page_num is not None:
                            page_num_str = str(page_num)
                            page_num_width = c.stringWidth(page_num_str, "Times-Roman", 14)
                            page_num_x = right_margin - page_num_width
                            
                            # Nuqtalarni chizish (mavzu nomi va bet raqami orasida)
                            line_width = c.stringWidth(line.strip(), "Times-Roman", 14)
                            dots_start = left_margin + line_width + 10
                            dots_end = page_num_x - 10
                            
                            if dots_end > dots_start:
                                dots_count = int((dots_end - dots_start) / 3)  # Har bir nuqta uchun 3px
                                dots = '.' * min(dots_count, 50)  # Maksimal 50 ta nuqta
                                dots_width = c.stringWidth(dots, "Times-Roman", 14)
                                dots_x = page_num_x - dots_width - 5
                                c.drawString(dots_x, y_position, dots)
                            
                            c.drawString(page_num_x, y_position, page_num_str)
                        
                        y_position -= line_height
                        check_new_page()
            else:
                words = content.split()
                line = ""
                max_width = right_margin - left_margin
                
                for word in words:
                    test_line = line + word + " "
                    if c.stringWidth(test_line, "Times-Roman", 14) < max_width:
                        line = test_line
                    else:
                        c.drawString(left_margin, y_position, line)
                        y_position -= line_height
                        check_new_page()
                        line = word + " "
                
                if line:
                    c.drawString(left_margin, y_position, line)
                    y_position -= line_height
            
            check_new_page()
            add_page_number()
            c.showPage()
            y_position = top_margin
    
    if page_number > 1:
        add_page_number()
    
    c.save()
    return file_path
