from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

def create_kurs_ishi_docx(content: dict, filename: str, topic: str, fish: str, university: str, subject: str, course_number: int) -> str:
    """Professional kurs ishi DOCX yaratish"""
    try:
        if not content or not content.get("full_content"):
            raise ValueError("Kontent bo'sh yoki noto'g'ri")
        
        os.makedirs("generated_files", exist_ok=True)
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
        
        _add_title_page(doc, topic, fish, university, subject, course_number)
        
        doc.add_page_break()
        
        _add_reja(doc)
        
        doc.add_page_break()
        
        if content.get("kirish"):
            _add_section_heading(doc, "KIRISH")
            _add_section_content(doc, content["kirish"])
            doc.add_page_break()
        
        if content.get("bob1"):
            _add_section_heading(doc, "I BOB. NAZARIY ASOSLAR")
            _add_section_content(doc, content["bob1"])
            doc.add_page_break()
        
        if content.get("bob2"):
            _add_section_heading(doc, "II BOB. AMALIY TAHLIL")
            _add_section_content(doc, content["bob2"])
            doc.add_page_break()
        
        if content.get("bob3"):
            _add_section_heading(doc, "III BOB. TAKLIFLAR VA YECHIMLAR")
            _add_section_content(doc, content["bob3"])
            doc.add_page_break()
        
        if content.get("xulosa"):
            _add_section_heading(doc, "XULOSA VA TAKLIFLAR")
            _add_section_content(doc, content["xulosa"])
            doc.add_page_break()
        
        if content.get("adabiyotlar"):
            _add_section_heading(doc, "FOYDALANILGAN ADABIYOTLAR RO'YXATI")
            _add_section_content(doc, content["adabiyotlar"])
            doc.add_page_break()
        
        if content.get("ilovalar"):
            _add_section_heading(doc, "ILOVALAR")
            _add_section_content(doc, content["ilovalar"])
            doc.add_page_break()
        
        _add_mundarija(doc)
        
        filepath = f"generated_files/{filename}.docx"
        doc.save(filepath)
        
        if not os.path.exists(filepath):
            raise Exception("Fayl yaratilmadi")
        
        logger.info(f"Professional kurs ishi DOCX yaratildi: {filepath}")
        return filepath
    
    except Exception as e:
        logger.error(f"DOCX yaratishda xatolik: {e}")
        raise Exception(f"DOCX fayl yaratishda xatolik: {str(e)}")

def _add_title_page(doc: Document, topic: str, fish: str, university: str, subject: str, course_number: int):
    """Titul varaq qo'shish"""
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("O'ZBEKISTON RESPUBLIKASI\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    run.font.all_caps = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{university.upper()}\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    run.font.all_caps = True
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KURS ISHI\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{subject} fanidan\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.italic = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Mavzu: {topic}\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = "Bajardi:"
    table.cell(0, 1).text = f"{fish}\n{course_number}-kurs talabasi"
    
    table.cell(1, 0).text = "Tekshirdi:"
    table.cell(1, 1).text = "_________________"
    
    table.cell(2, 0).text = "Topshirdi:"
    table.cell(2, 1).text = datetime.now().strftime("%d.%m.%Y")
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Toshkent – {datetime.now().year}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

def _add_reja(doc: Document):
    """Reja (Plan) qo'shish - titul varaqdan keyin"""
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("REJA")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    run.font.all_caps = True
    
    doc.add_paragraph()
    
    reja_items = [
        "KIRISH",
        "I BOB. NAZARIY ASOSLAR",
        "    1.1. Asosiy tushunchalar va ta'riflar",
        "    1.2. Nazariy yondashuvlar va konsepsiyalar",
        "    1.3. Xorijiy va mahalliy tajribalar tahlili",
        "    1.4. Ilmiy adabiyotlar sharhi",
        "II BOB. AMALIY TAHLIL",
        "    2.1. Joriy holat tahlili",
        "    2.2. Muammolar va ularning sabablari",
        "    2.3. Raqamli ma'lumotlar va statistik tahlil",
        "    2.4. Amaliy misollar va keis-tadqiqotlar",
        "III BOB. TAKLIFLAR VA YECHIMLAR",
        "    3.1. Muammolarni hal qilish yo'llari",
        "    3.2. Takomillashtirilgan yechimlar",
        "    3.3. Amalga oshirish mexanizmlari",
        "    3.4. Kutilayotgan natijalar va samaradorlik",
        "XULOSA VA TAKLIFLAR",
        "FOYDALANILGAN ADABIYOTLAR RO'YXATI",
        "ILOVALAR",
    ]
    
    for item in reja_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def _add_mundarija(doc: Document):
    """Mundarija qo'shish - oxirida, sahifa raqamlari bilan"""
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MUNDARIJA")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    run.font.all_caps = True
    
    doc.add_paragraph()
    
    mundarija_items = [
        ("KIRISH", "3"),
        ("I BOB. NAZARIY ASOSLAR", "7"),
        ("    1.1. Asosiy tushunchalar va ta'riflar", "7"),
        ("    1.2. Nazariy yondashuvlar va konsepsiyalar", "10"),
        ("    1.3. Xorijiy va mahalliy tajribalar tahlili", "12"),
        ("    1.4. Ilmiy adabiyotlar sharhi", "14"),
        ("II BOB. AMALIY TAHLIL", "16"),
        ("    2.1. Joriy holat tahlili", "16"),
        ("    2.2. Muammolar va ularning sabablari", "20"),
        ("    2.3. Raqamli ma'lumotlar va statistik tahlil", "22"),
        ("    2.4. Amaliy misollar va keis-tadqiqotlar", "25"),
        ("III BOB. TAKLIFLAR VA YECHIMLAR", "28"),
        ("    3.1. Muammolarni hal qilish yo'llari", "28"),
        ("    3.2. Takomillashtirilgan yechimlar", "30"),
        ("    3.3. Amalga oshirish mexanizmlari", "32"),
        ("    3.4. Kutilayotgan natijalar va samaradorlik", "34"),
        ("XULOSA VA TAKLIFLAR", "36"),
        ("FOYDALANILGAN ADABIYOTLAR RO'YXATI", "39"),
        ("ILOVALAR", "42"),
    ]
    
    for item, page in mundarija_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        p.add_run("." * (70 - len(item) - len(page)))
        
        run = p.add_run(page)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def _add_section_heading(doc: Document, heading_text: str):
    """Bo'lim sarlavhasini qo'shish"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(heading_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    run.font.all_caps = True
    
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_after = Pt(6)
    paragraph_format.space_before = Pt(12)

def _add_section_content(doc: Document, content: str):
    """Bo'lim kontentini qo'shish"""
    paragraphs = content.split('\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        run = p.add_run(para_text.strip())
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.first_line_indent = Cm(1.25)
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)

def _add_content(doc: Document, content: str):
    """Asosiy kontent qo'shish"""
    
    paragraphs = content.split('\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        
        p = doc.add_paragraph()
        
        is_heading = False
        if any(keyword in para_text.upper() for keyword in ['KIRISH', 'BOB', 'XULOSA', 'ADABIYOTLAR', 'I.', 'II.', 'III.']):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_heading = True
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        run = p.add_run(para_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        if is_heading:
            run.bold = True
        
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        if not is_heading:
            paragraph_format.first_line_indent = Cm(1.25)
        
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
